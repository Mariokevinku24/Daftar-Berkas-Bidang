import streamlit as st
from docx import Document
from io import BytesIO

st.title("📄 Merge Beberapa File Word (.docx) Menjadi Satu")

st.write("""
Unggah beberapa file Word yang ingin Anda gabungkan.
File hasil merge akan bisa diunduh.
""")

uploaded_files = st.file_uploader("Upload file .docx", type=["docx"], accept_multiple_files=True)

if uploaded_files:
    if st.button("Gabungkan File"):
        merged_document = Document()

        for index, uploaded_file in enumerate(uploaded_files):
            doc = Document(uploaded_file)
            # Tambahkan setiap paragraf dari file
            for paragraph in doc.paragraphs:
                merged_document.add_paragraph(paragraph.text)
            
            # Tambahkan page break antar dokumen (opsional)
            if index < len(uploaded_files) - 1:
                merged_document.add_page_break()

        # Simpan ke buffer
        buffer = BytesIO()
        merged_document.save(buffer)
        buffer.seek(0)

        st.success("File berhasil digabungkan!")

        st.download_button(
            label="📥 Download File Word Gabungan",
            data=buffer,
            file_name="merged_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
